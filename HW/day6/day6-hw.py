import requests
import pandas as pd
import uuid
import time
from docx import Document
from qdrant_client import QdrantClient, models
from openai import OpenAI
from deepeval.metrics import FaithfulnessMetric
from deepeval.test_case import LLMTestCase
from deepeval.models.base_model import DeepEvalBaseLLM

# --- 1. 配置與初始化 ---
QDRANT_URL = "http://localhost:6333"
LLM_BASE_URL = "https://ws-03.wade0426.me/v1"
EMBED_API_URL = "https://ws-04.wade0426.me/embed"
MODEL_NAME = "/models/gpt-oss-120b"
COLLECTION_NAME = "water_qa_hw"

client = QdrantClient(url=QDRANT_URL)
llm_client = OpenAI(base_url=LLM_BASE_URL, api_key="no-key")

# 自定義 DeepEval 模型類別
class CustomLLM(DeepEvalBaseLLM):
    def __init__(self, model_name): self.model_name = model_name
    def load_model(self): return llm_client
    def generate(self, prompt: str) -> str:
        res = self.load_model().chat.completions.create(model=self.model_name, messages=[{"role": "user", "content": prompt}])
        return res.choices[0].message.content
    async def a_generate(self, prompt: str) -> str: return self.generate(prompt)
    def get_model_name(self): return self.model_name

custom_eval_model = CustomLLM(model_name=MODEL_NAME)

# --- 2. 工具函數 ---

def get_embeddings(texts):
    res = requests.post(EMBED_API_URL, json={"texts": texts, "task_description": "檢索台水文件", "normalize": True})
    return res.json()["embeddings"]

def hybrid_search(query_text):
    vector = get_embeddings([query_text])[0]
    search_result = client.query_points(
        collection_name=COLLECTION_NAME,
        prefetch=[
            models.Prefetch(query=vector, using="dense", limit=5),
            models.Prefetch(query=models.SparseVector(indices=[1], values=[1.0]), using="sparse", limit=5),
        ],
        query=models.FusionQuery(fusion=models.Fusion.RRF),
        limit=3
    )
    return [hit.payload['text'] for hit in search_result.points]

# --- 3. 執行流程 ---

if __name__ == "__main__":
    # --- A. 資料匯入 (含分批邏輯解決 400 錯誤) ---
    print("Step 1: 正在讀取並匯入資料...")
    doc = Document("qa_data.docx")
    paragraphs = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 10]
    
    # 取得第一筆確認維度
    first_embedding = get_embeddings(paragraphs[:1])[0]
    vector_dim = len(first_embedding)

    if client.collection_exists(COLLECTION_NAME):
        client.delete_collection(COLLECTION_NAME)
    
    client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config={"dense": models.VectorParams(size=vector_dim, distance=models.Distance.COSINE)},
        sparse_vectors_config={"sparse": models.SparseVectorParams(modifier=models.Modifier.IDF)}
    )

    # 分批轉換與寫入 (每 100 筆一組，避免 Payload 太大)
    chunk_size = 100
    for i in range(0, len(paragraphs), chunk_size):
        batch_texts = paragraphs[i:i+chunk_size]
        batch_vecs = get_embeddings(batch_texts)
        
        points = [
            models.PointStruct(
                id=str(uuid.uuid4()),
                vector={"dense": v, "sparse": models.Document(text=t, model="Qdrant/bm25")},
                payload={"text": t}
            ) for t, v in zip(batch_texts, batch_vecs)
        ]
        client.upsert(collection_name=COLLECTION_NAME, points=points)
        print(f"匯入進度: {min(i+chunk_size, len(paragraphs))}/{len(paragraphs)}")

    # --- B. 執行前三題驗證 ---
    print("\nStep 2: 開始執行前三題評測...")
    df = pd.read_csv("day6_HW_questions.csv.xlsx - day6_HW_questions.csv.csv")
    
    for i in range(3):
        question = df.loc[i, 'questions']
        print(f"處理中: {question}")
        
        # 檢索與回答
        contexts = hybrid_search(question)
        prompt = f"你是台水客服，請僅根據以下內容回答問題：\n{chr(10).join(contexts)}\n\n問題：{question}"
        ans = llm_client.chat.completions.create(model=MODEL_NAME, messages=[{"role": "user", "content": prompt}]).choices[0].message.content
        
        # 評測
        test_case = LLMTestCase(input=question, actual_output=ans, retrieval_context=contexts)
        metric = FaithfulnessMetric(model=custom_eval_model)
        metric.measure(test_case)
        
        df.at[i, 'answer'] = ans
        df.at[i, 'Faithfulness'] = metric.score
        print(f"完成！分數: {metric.score}")

    # 儲存結果
    df.head(3).to_csv("day6_results_top3.csv", index=False, encoding='utf-8-sig')
    print("\n全部流程完成！請檢查 day6_results_top3.csv")